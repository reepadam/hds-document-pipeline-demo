"""
HDS Demo — Pipeline module
OCR + LLM extraction logic, decoupled from Streamlit UI for testability.
"""
import json
import os
import re
from io import BytesIO
from pathlib import Path

from google.cloud import vision
import anthropic
from PIL import Image
import pypdfium2 as pdfium
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPM


DEFAULT_CREDS_PATH = r"C:\Users\Adam\Documents\Claude\Job\cloudvision-497116-3258213b51c0.json"
DEFAULT_ANTHROPIC_KEY_PATH = r"C:\Users\Adam\Documents\Claude\Job\Anthropic_API.txt"

CREDS_PATH = os.environ.get("HDS_DEMO_GCP_CREDS", DEFAULT_CREDS_PATH)
ANTHROPIC_KEY_PATH = os.environ.get("HDS_DEMO_ANTHROPIC_KEY_PATH", DEFAULT_ANTHROPIC_KEY_PATH)
EXTRACTION_MODEL = "claude-haiku-4-5"


# When running on Streamlit Community Cloud, credentials come from st.secrets
# (not from disk). Try to load them at import time; fall back to local files
# if running on Adam's Windows machine.
_anthropic_api_key_from_secrets = None
try:
    import streamlit as _st
    if hasattr(_st, "secrets"):
        # GCP service account JSON content under [gcp_service_account]
        if "gcp_service_account" in _st.secrets:
            import json as _json
            import tempfile as _tempfile
            gcp_sa = dict(_st.secrets["gcp_service_account"])
            _tf = _tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False)
            _json.dump(gcp_sa, _tf)
            _tf.close()
            os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = _tf.name
            CREDS_PATH = _tf.name
        if "anthropic_api_key" in _st.secrets:
            _anthropic_api_key_from_secrets = str(_st.secrets["anthropic_api_key"])
except Exception:
    pass

# Fallback: local file for GCP creds (Adam's dev machine)
if "GOOGLE_APPLICATION_CREDENTIALS" not in os.environ:
    if os.path.exists(CREDS_PATH):
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = CREDS_PATH


# Wholesale blank-garment cost reference (HDS-typical 2026 distributor pricing).
# Used by the order builder to roll a full per-piece quote.
BASE_GARMENT_COST_USD = {
    "T-shirt (100% cotton)": 4.25,
    "Polo shirt (cotton-poly blend)": 11.00,
    "Hoodie / sweatshirt (heavy fleece)": 18.50,
    "Cap / hat (structured)": 6.75,
    "Canvas tote bag": 3.50,
    "Performance / athletic (poly)": 9.25,
    "Denim jacket": 28.00,
    "Towel (terrycloth)": 7.50,
}

# Garment-appropriate color options. Curated to what actually exists in
# blank-apparel catalogs (SanMar, S&S, Alphabroder).
GARMENT_COLORS = {
    "T-shirt (100% cotton)": [
        "White", "Black", "Navy", "Royal Blue", "Carolina Blue", "Red",
        "Maroon", "Forest Green", "Kelly Green", "Sport Gray", "Charcoal",
        "Gold", "Orange", "Purple", "Pink",
    ],
    "Polo shirt (cotton-poly blend)": [
        "White", "Black", "Navy", "Red", "Royal Blue", "Forest Green",
        "Burgundy", "Heather Gray", "Light Blue", "Khaki", "Stone", "Hunter Green",
    ],
    "Hoodie / sweatshirt (heavy fleece)": [
        "Black", "Navy", "Charcoal", "Sport Gray", "Heather Gray", "Maroon",
        "Forest Green", "Royal Blue", "White", "Red", "Athletic Gold", "Burgundy",
    ],
    "Cap / hat (structured)": [
        "Black", "Navy", "Royal Blue", "Red", "White", "Khaki", "Charcoal",
        "Stone", "Hunter Green", "Heather Gray", "Pink", "Camo",
    ],
    "Canvas tote bag": [
        "Natural", "Black", "Navy", "Red", "Royal Blue", "Forest Green",
        "Yellow", "Pink",
    ],
    "Performance / athletic (poly)": [
        "White", "Black", "Navy", "Red", "Royal Blue", "Sport Gray",
        "Athletic Gold", "Carolina Blue", "Forest Green", "Purple", "Hot Pink",
    ],
    "Denim jacket": [
        "Light Stone Wash", "Medium Wash", "Dark Indigo", "Black Denim", "Vintage Wash",
    ],
    "Towel (terrycloth)": [
        "White", "Black", "Navy", "Red", "Royal Blue", "Hunter Green",
        "Maroon", "Khaki", "Charcoal", "Light Blue",
    ],
}

# Size options per garment. Apparel uses standard XS-4XL, caps use fit
# sizes, totes and towels use dimension-based sizes.
GARMENT_SIZES = {
    "T-shirt (100% cotton)": ["XS", "S", "M", "L", "XL", "2XL", "3XL", "4XL"],
    "Polo shirt (cotton-poly blend)": ["XS", "S", "M", "L", "XL", "2XL", "3XL"],
    "Hoodie / sweatshirt (heavy fleece)": ["S", "M", "L", "XL", "2XL", "3XL"],
    "Performance / athletic (poly)": ["XS", "S", "M", "L", "XL", "2XL", "3XL"],
    "Denim jacket": ["S", "M", "L", "XL", "2XL", "3XL"],
    "Cap / hat (structured)": ["OSFA", "S/M", "L/XL"],
    "Canvas tote bag": ["One Size"],
    "Towel (terrycloth)": ["Hand (11x18)", "Sport (16x27)", "Beach (30x60)"],
}

# Logo placement options per garment, with default size in inches (width, height).
# Cost scales with area - a 2x1 wrist hit costs far less than a 12x14 back hit.
# Format: { garment: { placement_name: (default_w_in, default_h_in) } }
GARMENT_PLACEMENTS = {
    "T-shirt (100% cotton)": {
        "Left Chest": (3.5, 3.0),
        "Right Chest": (3.5, 3.0),
        "Center Chest": (10.0, 8.0),
        "Full Front": (12.0, 12.0),
        "Full Back": (12.0, 14.0),
        "Back Yoke (upper)": (10.0, 4.0),
        "Left Sleeve": (3.0, 3.0),
        "Right Sleeve": (3.0, 3.0),
    },
    "Polo shirt (cotton-poly blend)": {
        "Left Chest": (3.5, 3.0),
        "Right Chest": (3.5, 3.0),
        "Left Sleeve": (3.0, 3.0),
        "Right Sleeve": (3.0, 3.0),
        "Back Yoke (upper)": (10.0, 4.0),
    },
    "Hoodie / sweatshirt (heavy fleece)": {
        "Left Chest": (3.5, 3.0),
        "Right Chest": (3.5, 3.0),
        "Center Chest": (10.0, 8.0),
        "Full Front": (12.0, 12.0),
        "Full Back": (12.0, 14.0),
        "Hood (left)": (2.5, 2.0),
        "Left Sleeve": (3.0, 3.0),
        "Right Sleeve": (3.0, 3.0),
    },
    "Performance / athletic (poly)": {
        "Left Chest": (3.5, 3.0),
        "Right Chest": (3.5, 3.0),
        "Center Chest": (10.0, 8.0),
        "Full Front": (12.0, 12.0),
        "Full Back": (12.0, 14.0),
        "Left Sleeve": (3.0, 3.0),
        "Right Sleeve": (3.0, 3.0),
    },
    "Denim jacket": {
        "Left Chest": (3.5, 3.0),
        "Right Chest": (3.5, 3.0),
        "Full Back": (10.0, 12.0),
        "Left Sleeve": (3.0, 3.0),
        "Right Sleeve": (3.0, 3.0),
    },
    "Cap / hat (structured)": {
        "Front Center": (4.0, 2.25),
        "Side Left": (2.5, 1.5),
        "Side Right": (2.5, 1.5),
        "Back": (3.5, 1.5),
    },
    "Canvas tote bag": {
        "Center Front": (8.0, 8.0),
        "Center Back": (8.0, 8.0),
    },
    "Towel (terrycloth)": {
        "Corner Embroidery": (3.0, 3.0),
        "Center Hem": (6.0, 4.0),
    },
}

# Reference: per-piece costs are anchored to a "typical left chest" sizing.
# We use 10 sq in as the baseline (~3.5 x 2.8 in).
BASELINE_LOGO_AREA_SQIN = 10.0

# Decoration-method cost reference (HDS-typical 2026 industry pricing).
# Used by build_quote_line so each method has its own materials/labor/setup
# regardless of which method the LLM happened to recommend.
METHOD_COSTS = {
    "embroidery": {
        "label": "Embroidery",
        "materials_per_pc": 0.50,    # thread + backing + bobbin
        "labor_per_pc": 1.20,        # machine time + handling
        "setup_one_time": 32.00,     # digitizing fee (one-time per design)
    },
    "screen_print": {
        "label": "Screen print",
        "materials_per_pc": 0.20,    # plastisol ink
        "labor_per_pc": 0.60,        # press operator time
        "setup_one_time": 25.00,     # screen burn (per color, simplified to one)
    },
    "dtg": {
        "label": "DTG (direct-to-garment)",
        "materials_per_pc": 0.80,    # DTG ink + pretreatment
        "labor_per_pc": 0.50,        # short cycle time
        "setup_one_time": 0.00,      # no real setup
    },
    "heat_transfer": {
        "label": "Heat transfer / vinyl",
        "materials_per_pc": 0.40,
        "labor_per_pc": 0.70,
        "setup_one_time": 5.00,
    },
    "sublimation": {
        "label": "Sublimation (poly garments only)",
        "materials_per_pc": 0.30,
        "labor_per_pc": 0.50,
        "setup_one_time": 5.00,
    },
}

# Smart default decoration method per garment. Order builder uses this when
# the user adds a new line - they can override via the per-line method dropdown.
GARMENT_METHOD_DEFAULTS = {
    "T-shirt (100% cotton)": "screen_print",
    "Polo shirt (cotton-poly blend)": "embroidery",
    "Hoodie / sweatshirt (heavy fleece)": "embroidery",
    "Cap / hat (structured)": "embroidery",
    "Canvas tote bag": "screen_print",
    "Performance / athletic (poly)": "sublimation",
    "Denim jacket": "embroidery",
    "Towel (terrycloth)": "embroidery",
}


def fmt_money(value, blank="—"):
    """Format any amount-like value as $X.XX. Handles strings like '165.0',
    '$165', '165', floats, and ints. Returns blank string if no amount."""
    if value is None or value == "" or value == "—":
        return blank
    try:
        s = str(value).replace("$", "").replace(",", "").strip()
        if not s:
            return blank
        return f"${float(s):,.2f}"
    except (ValueError, TypeError):
        return str(value) if value else blank


def scale_decoration_cost(base_cost_per_pc, logo_area_sqin, setup_per_pc=None):
    """Scale the LLM's baseline decoration cost by actual logo area.

    Decoration cost has two scalable components:
    - materials + labor (scale roughly linearly with area)
    - setup (digitizing fee or screen burn cost) - scales more slowly,
      use sqrt to reflect that bigger designs need somewhat more setup
      effort but not proportionally more

    base_cost_per_pc is the LLM's reported decoration cost at baseline area.
    Returns scaled cost per piece.
    """
    if base_cost_per_pc is None:
        return 0.0
    area_ratio = max(logo_area_sqin, 0.5) / BASELINE_LOGO_AREA_SQIN
    if setup_per_pc is None:
        # Treat the entire decoration cost as linearly scaling by area
        return base_cost_per_pc * area_ratio
    variable_cost = base_cost_per_pc - setup_per_pc
    scaled_variable = max(variable_cost, 0.0) * area_ratio
    scaled_setup = setup_per_pc * (area_ratio ** 0.5)
    return scaled_variable + scaled_setup


def _load_anthropic_key():
    # Prefer Streamlit secrets (production)
    if _anthropic_api_key_from_secrets:
        return _anthropic_api_key_from_secrets
    # Fall back to local file (Adam's dev machine)
    with open(ANTHROPIC_KEY_PATH) as f:
        return f.read().strip()


_vision_client = None
_anthropic_client = None


def get_vision_client():
    global _vision_client
    if _vision_client is None:
        _vision_client = vision.ImageAnnotatorClient()
    return _vision_client


def get_anthropic_client():
    global _anthropic_client
    if _anthropic_client is None:
        _anthropic_client = anthropic.Anthropic(api_key=_load_anthropic_key())
    return _anthropic_client


def pdf_to_images(pdf_bytes):
    pdf = pdfium.PdfDocument(BytesIO(pdf_bytes))
    return [page.render(scale=2.0).to_pil() for page in pdf]


def svg_to_pil(svg_bytes):
    """Rasterize an SVG (bytes) into a PIL image. Vision API can't read SVG directly."""
    drawing = svg2rlg(BytesIO(svg_bytes))
    if drawing is None:
        raise RuntimeError("svg2rlg returned None - could not parse SVG")
    if drawing.width and drawing.height:
        target_max = 1024
        scale = target_max / max(drawing.width, drawing.height)
        if scale > 1:
            drawing.width *= scale
            drawing.height *= scale
            drawing.scale(scale, scale)
    png_bytes = renderPM.drawToString(drawing, fmt="PNG")
    return Image.open(BytesIO(png_bytes))


def image_to_bytes(img, fmt="PNG"):
    buf = BytesIO()
    if img.mode == "RGBA" and fmt == "JPEG":
        img = img.convert("RGB")
    img.save(buf, format=fmt)
    return buf.getvalue()


def ocr_image_bytes(image_bytes):
    client = get_vision_client()
    image = vision.Image(content=image_bytes)
    response = client.document_text_detection(image=image)
    if response.error.message:
        raise RuntimeError(f"Cloud Vision error: {response.error.message}")
    return response.full_text_annotation.text or ""


def analyze_image_features(image_bytes):
    """Multi-feature Vision call: labels, logos, colors, text."""
    client = get_vision_client()
    image = vision.Image(content=image_bytes)
    features = [
        vision.Feature(type_=vision.Feature.Type.IMAGE_PROPERTIES),
        vision.Feature(type_=vision.Feature.Type.LABEL_DETECTION, max_results=10),
        vision.Feature(type_=vision.Feature.Type.LOGO_DETECTION),
        vision.Feature(type_=vision.Feature.Type.TEXT_DETECTION),
    ]
    response = client.annotate_image(
        vision.AnnotateImageRequest(image=image, features=features)
    )
    colors = []
    if response.image_properties_annotation and response.image_properties_annotation.dominant_colors:
        for c in response.image_properties_annotation.dominant_colors.colors[:5]:
            colors.append({
                "rgb": {"r": int(c.color.red), "g": int(c.color.green), "b": int(c.color.blue)},
                "hex": f"#{int(c.color.red):02x}{int(c.color.green):02x}{int(c.color.blue):02x}",
                "score": round(c.score, 3),
                "pixel_fraction": round(c.pixel_fraction, 3),
            })
    labels = [{"description": l.description, "score": round(l.score, 3)} for l in response.label_annotations[:10]]
    detected_logos = [{"description": l.description, "score": round(l.score, 3)} for l in response.logo_annotations]
    embedded_text = (response.text_annotations[0].description if response.text_annotations else "").strip()
    return {
        "dominant_colors": colors,
        "detected_labels": labels,
        "detected_logos": detected_logos,
        "embedded_text": embedded_text,
    }


def auto_detect_mode(image_bytes, file_extension):
    """Return ('logo'|'document', features_dict_or_None) based on file type and Vision analysis."""
    ext = file_extension.lower().lstrip(".")
    if ext == "svg":
        return "logo", None
    features = analyze_image_features(image_bytes)
    text_len = len(features["embedded_text"])
    label_words = " ".join(l["description"].lower() for l in features["detected_labels"])
    doc_signals = ["document", "receipt", "invoice", "paper", "text", "menu", "letter", "envelope", "page"]
    art_signals = ["logo", "graphics", "graphic design", "brand", "emblem", "symbol", "trademark", "icon", "illustration", "artwork", "drawing", "vector"]
    doc_score = sum(1 for s in doc_signals if s in label_words)
    art_score = sum(1 for s in art_signals if s in label_words)
    if text_len > 150 and doc_score >= art_score:
        return "document", features
    if text_len < 30 and art_score >= 1:
        return "logo", features
    if doc_score > art_score:
        return "document", features
    if art_score > doc_score:
        return "logo", features
    return ("document" if text_len > 80 else "logo"), features


def strip_json_fence(s):
    s = s.strip()
    s = re.sub(r"^```(?:json)?\s*", "", s)
    s = re.sub(r"\s*```$", "", s)
    return s


CLASSIFY_AND_EXTRACT_PROMPT = """You are a document-extraction assistant for a promotional products distributor (HDS Marketing).

You will receive OCR text from a document. Your job is to:
1. Classify the document type (one of: vendor_invoice, freight_invoice, product_data_sheet, handwritten_note, mixed_form, order_request, receipt, unknown)
2. Extract structured data appropriate for that type
3. Score your confidence (high/medium/low)
4. Flag any fields where OCR may have introduced errors

OCR text:
---
{ocr_text}
---

Return ONLY valid JSON (no preamble, no markdown fence) with shape:

{{
  "document_type": "string",
  "confidence": "high|medium|low",
  "extracted_data": {{}},
  "flags": [],
  "notes": "string"
}}

For vendor_invoice: include vendor (name, address, phone), transaction (date, invoice_number, po_number), bill_to, line_items (array of description, sku, quantity, unit_price, extended_price), totals (subtotal, tax, freight, total), payment (method, last_four).

For freight_invoice (UPS/FedEx/R+L/freight carriers): include carrier, tracking_number, ship_date, origin, destination, weight, service_level, charges (base, fuel_surcharge, accessorial, total), invoice_number, bill_to.

For product_data_sheet (Madeira thread, Wilflex ink, blank apparel specs): include manufacturer, product_name, product_line, product_type, composition, key_specs, test_ratings, mechanical_properties, care_instructions, certifications, document_date, manufacturer_contact.

For handwritten_note: include date, author, subject, key_fields, action_items, unresolved.

For mixed_form: include form_type, printed_fields, handwritten_fields, unfilled_fields.

For order_request: include customer, items_requested, decoration_spec, in_hands_date, special_instructions.

For receipt (point-of-sale receipt - meals, fuel, supplies, parking, etc.; usually <$200, has a time stamp, no PO number): include merchant (name, address, phone), date, time, line_items (array of description, quantity, price), subtotal, tax, tip (if present), total, payment_method (cash/credit/debit), card_last_four (if visible), and suggested_expense_category (one of: meals_dining, fuel, supplies, materials, parking, lodging, shipping, software_subscriptions, equipment, repairs_maintenance, other).

CRITICAL EXTRACTION POLICY: Extract EVERY field that has any data visible on the document, even partial. Do not leave a field blank just because the document looks atypical, isn't a perfect match for the document type, or appears to be from a non-promotional-products vendor. If the doc looks like a receipt but is actually a vendor invoice (or vice versa), classify as the better match AND still extract every available field using whichever schema fits. Treat the schema as a guide, not a gate. A SaaS subscription invoice, software receipt, or any unusual document should still surface its merchant, dates, amounts, line items, and payment info into the structured output. Aggressive extraction beats blank fields - the user can always correct, but cannot fill what wasn't extracted.

For dollar amounts always output as strings with two decimal places (e.g. "165.00" not "165" or "165.0"). Include the dollar amount without currency symbol.

Flag any field where OCR seems unclear. Notes field for anything that didn't fit the schema."""


LOGO_RECOMMENDATIONS_PROMPT = """You are a decoration-planning assistant for HDS Marketing - a promotional products distributor whose production facility (X-Deco, Bridgeville PA) runs 62 Tajima embroidery heads and multiple multi-color screen-printing presses.

A customer has provided artwork they want decorated on a garment. Based on the Vision analysis, the target garment, and the source format, recommend the optimal decoration approach.

ARTWORK ANALYSIS (from Google Cloud Vision):
- Dominant colors: {colors}
- Detected labels: {labels}
- Detected brand logos: {logos}
- Embedded text: {embedded_text}
- Source format: {input_format}   (svg = vector, clean digitizable art; png/jpg = raster, may need redraw)

TARGET GARMENT: {garment}

DECORATION-METHOD DECISION RULES (apply these strictly - do NOT default to screen print for everything):

Garment-driven defaults (override only with a stated reason):
- Cap / hat (structured) -> embroidery (3D foam embroidery for raised logos; flat embroidery otherwise). Screen print is rare on caps.
- Polo shirt -> embroidery on left chest. This is the industry standard for corporate polos.
- Hoodie / sweatshirt -> embroidery for left chest or small front logos; screen print only for large back/front-of-chest designs (>8 inches) or photographic art.
- Denim jacket -> embroidery (durable, classic). Patches (embroidered or woven) are also valid.
- Towel (terrycloth) -> embroidery only. Screen print and DTG do not adhere to terrycloth pile.
- Canvas tote bag -> screen print (large flat print area, low color count favors it).
- T-shirt (100% cotton) -> screen print for >=4 inches print size with <=6 spot colors; embroidery for left-chest corporate looks; DTG for photographic / >6 colors / very low quantities.
- Performance / athletic (poly) -> sublimation if the garment is white/light poly; otherwise plastisol screen print with a low-bleed underbase. Embroidery is acceptable for small left-chest only (puckers on thin poly).

Source-format signal:
- SVG (vector) is IDEAL for embroidery - clean paths digitize cleanly into stitch files with minimal manual cleanup. Treat SVG as a positive signal toward embroidery for any garment where embroidery is plausible.
- Raster (PNG/JPG) at low resolution or with photographic shading favors DTG or screen print; flag if vector redraw is needed before embroidery.

Color-count signal:
- <=3 spot colors -> screen print and embroidery both viable.
- 4-8 spot colors with flat fills -> screen print viable; embroidery viable but thread-change overhead grows.
- >8 colors or photographic gradients -> DTG or sublimation (NOT screen print, NOT embroidery).

Return ONLY valid JSON (no preamble, no markdown fence) with this exact shape. EVERY numeric field below MUST be populated with a realistic dollar/integer value - do NOT return null for any cost field. If you are uncertain, give your best estimate and note the assumption in cost_basis_notes.

{{
  "recommended_decoration_method": "embroidery|screen_print|dtg|heat_transfer|sublimation",
  "method_reasoning": "2-3 sentences. State WHY this method beats the alternatives for THIS garment and THIS artwork.",
  "thread_or_ink_recommendation": {{
    "brand_product": "e.g., Madeira Classic No. 40 for embroidery, or Wilflex Epic plastisol for screen print",
    "color_count": <integer>,
    "pantone_matches_needed": ["#hex1", "#hex2"]
  }},
  "machine_settings": {{
    "needle_size": "e.g., 75/11 sharp for embroidery, or null for non-embroidery",
    "mesh_count": "e.g., 156 for screen print, or null for non-screen-print",
    "stitch_density_or_ink_deposit": "string",
    "stabilizer_or_pretreatment": "string"
  }},
  "estimated_complexity": {{
    "stitch_count_estimate": <integer - required for embroidery, else null>,
    "color_separation_count": <integer - required for screen print, else null>,
    "estimated_run_time_per_piece_seconds": <integer, required>,
    "setup_time_minutes": <integer, required>
  }},
  "estimated_unit_cost": {{
    "materials_per_piece_usd": <number, required, e.g. 0.45>,
    "labor_per_piece_usd": <number, required, e.g. 1.20>,
    "setup_amortized_over_100pcs_usd": <number, required, e.g. 0.55>,
    "total_at_100pcs_usd": <number, required - sum of the three above>,
    "total_at_500pcs_usd": <number, required - same materials+labor, setup amortized over 500>,
    "cost_basis_notes": "1 sentence on assumptions (e.g. 'Assumes $0.0008/1000-stitch thread cost, $45/hr loaded labor, $25 digitizing setup')"
  }},
  "quality_flags": ["any concerns about the artwork - small text, thin lines, gradient, low resolution, trademark"],
  "next_steps_for_production": ["list of 2-4 specific actions production should take"],
  "production_sheet": {{
    "design_info": {{
      "size_inches": "e.g. '3.50 x 2.80' - typical left-chest is 3.5-4.0 inches wide",
      "stitch_count_total": <integer, e.g. 8500>,
      "trims": <integer, ~3x color_changes>,
      "color_changes": <integer>,
      "total_top_thread_meters": <number, ~stitch_count * 0.0045>,
      "total_bobbin_meters": <number, ~stitch_count * 0.0009>,
      "hoop": "e.g. 'Round (4.25 in.)' or 'Rectangle (5.5 x 9.4 in.)' - pick smallest hoop that fits the design"
    }},
    "active_colors": [
      {{
        "thread_code": "Madeira code, e.g. '1800'",
        "thread_name": "Madeira name, e.g. 'White'",
        "brand_line": "e.g. 'Madeira Classic 40' or 'Madeira Polyneon 40' or 'Madeira CR Metallic'",
        "hex": "#rrggbb - match to dominant color in artwork",
        "stitches": <integer - this thread's total stitches across the design>,
        "thread_meters": <number>
      }}
    ],
    "color_sequence": [
      {{
        "seq": <integer starting at 1>,
        "thread_code": "Madeira code",
        "thread_name": "Madeira name",
        "brand_line": "e.g. 'Madeira Classic 40'",
        "hex": "#rrggbb",
        "stitches": <integer for this sequence step>,
        "thread_meters": <number>
      }}
    ]
  }}
}}

PRODUCTION SHEET RULES (REQUIRED when method == embroidery, set production_sheet to null otherwise):
- Map each dominant color in the artwork to a real Madeira thread. Common Madeira product lines: 'Madeira Classic 40' (rayon, default), 'Madeira Polyneon 40' (poly, for safety/medical/bleach-resistant), 'Madeira CR Metallic' (sparkle/metallic), 'Madeira Frosted Matt' (matte finish), 'Madeira Supertwist' (specialty).
- Use realistic Madeira thread codes (e.g. 1800 White, 1801 Super White, 1810 Black, 1761 Cardinal Red, 1133 Navy, 1101 Gold, 1410 Royal Blue). It's OK to invent plausible codes if you don't know exact ones - note that in cost_basis_notes.
- 'active_colors' lists unique threads (deduplicated). 'color_sequence' lists every stitching step in order (same thread can appear multiple times if the design changes back to it).
- For a typical left-chest logo expect 6,000-12,000 total stitches. For a back-of-jacket logo expect 25,000-60,000.
- Sum of stitches in active_colors should equal stitch_count_total. Sum of stitches in color_sequence should also equal stitch_count_total.

Use realistic 2026 promotional-products industry pricing. Be specific to the artwork - don't give generic advice."""


def extract_structured(ocr_text):
    client = get_anthropic_client()
    prompt = CLASSIFY_AND_EXTRACT_PROMPT.format(ocr_text=ocr_text)
    msg = client.messages.create(
        model=EXTRACTION_MODEL, max_tokens=2500,
        messages=[{"role": "user", "content": prompt}],
    )
    parsed = json.loads(strip_json_fence(msg.content[0].text))
    parsed["_meta"] = {
        "model": msg.model,
        "tokens_in": msg.usage.input_tokens,
        "tokens_out": msg.usage.output_tokens,
        "approx_cost_usd": round((msg.usage.input_tokens * 0.0000008) + (msg.usage.output_tokens * 0.000004), 4),
    }
    return parsed


def get_logo_recommendations(vision_features, garment, input_format="png"):
    client = get_anthropic_client()
    prompt = LOGO_RECOMMENDATIONS_PROMPT.format(
        colors=json.dumps(vision_features["dominant_colors"]),
        labels=json.dumps(vision_features["detected_labels"]),
        logos=json.dumps(vision_features["detected_logos"]),
        embedded_text=vision_features["embedded_text"] or "(none)",
        garment=garment,
        input_format=input_format,
    )
    msg = client.messages.create(
        model=EXTRACTION_MODEL, max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )
    parsed = json.loads(strip_json_fence(msg.content[0].text))
    parsed["_meta"] = {
        "model": msg.model,
        "tokens_in": msg.usage.input_tokens,
        "tokens_out": msg.usage.output_tokens,
        "approx_cost_usd": round((msg.usage.input_tokens * 0.0000008) + (msg.usage.output_tokens * 0.000004), 4),
    }
    return parsed


def build_quote_line(garment_type, base_color, quantity, recs, markup_pct=35.0,
                     logo_width_in=None, logo_height_in=None, placement=None, method=None):
    """Roll a single line item into a full quote: blank + (method-driven, size-scaled) decoration + markup.

    method: one of METHOD_COSTS keys (embroidery, screen_print, dtg, heat_transfer,
    sublimation). If None, defaults to embroidery. Each method has its own
    materials/labor cost from the METHOD_COSTS table - the LLM's recommendation
    is just a smart default; the actual cost basis is industry-standard per-method
    pricing so users can mix methods across lines without re-prompting Claude.

    Logo size scales the variable per-piece costs by area relative to baseline.
    """
    blank = BASE_GARMENT_COST_USD.get(garment_type, 0.0)

    # Decoration-method-driven per-piece costs (industry standard table).
    # Method is now an explicit per-line choice, not derived from the LLM's
    # single recommendation - the user can mix methods across lines.
    method_key = method or "embroidery"
    method_data = METHOD_COSTS.get(method_key, METHOD_COSTS["embroidery"])
    materials = method_data["materials_per_pc"]
    labor = method_data["labor_per_pc"]
    baseline_variable = materials + labor

    # Scale variable costs by logo area
    if logo_width_in and logo_height_in and logo_width_in > 0 and logo_height_in > 0:
        logo_area = logo_width_in * logo_height_in
        area_ratio = logo_area / BASELINE_LOGO_AREA_SQIN
        decoration_cost = baseline_variable * max(area_ratio, 0.05)
    else:
        decoration_cost = baseline_variable
        logo_area = BASELINE_LOGO_AREA_SQIN
        area_ratio = 1.0

    house_cost_per_pc = blank + decoration_cost
    customer_price_per_pc = house_cost_per_pc * (1 + markup_pct / 100.0)
    line_total = customer_price_per_pc * quantity

    return {
        "garment_type": garment_type,
        "base_color": base_color,
        "quantity": quantity,
        "placement": placement,
        "method": method_key,
        "method_label": method_data["label"],
        "logo_width_in": logo_width_in,
        "logo_height_in": logo_height_in,
        "logo_area_sqin": round(logo_area, 2),
        "size_scale_factor": round(area_ratio, 2),
        "blank_cost_per_pc": round(blank, 2),
        "materials_per_pc": round(materials * max(area_ratio, 0.05), 2),
        "labor_per_pc": round(labor * max(area_ratio, 0.05), 2),
        "decoration_cost_per_pc": round(decoration_cost, 2),
        "house_cost_per_pc": round(house_cost_per_pc, 2),
        "customer_price_per_pc": round(customer_price_per_pc, 2),
        "line_total": round(line_total, 2),
        "markup_pct": markup_pct,
    }


def get_order_setup_fee(lines, markup_pct=35.0):
    """One-time setup fee summed across UNIQUE methods used in the order.

    Each decoration method needs its own setup (embroidery digitizing,
    screen burning, etc.). If an order has both embroidered caps AND
    screen-printed t-shirts, the customer pays setup for both methods -
    once each, regardless of how many units per method.
    """
    methods_used = set()
    for ln in (lines or []):
        m = ln.get("method")
        if m and m in METHOD_COSTS:
            methods_used.add(m)
    if not methods_used:
        methods_used = {"embroidery"}  # fallback

    breakdown = []
    for m in sorted(methods_used):
        breakdown.append({
            "method": m,
            "method_label": METHOD_COSTS[m]["label"],
            "setup_house": METHOD_COSTS[m]["setup_one_time"],
            "setup_customer": round(METHOD_COSTS[m]["setup_one_time"] * (1 + markup_pct / 100.0), 2),
        })
    setup_house = sum(b["setup_house"] for b in breakdown)
    setup_customer = sum(b["setup_customer"] for b in breakdown)
    return {
        "setup_house_usd": round(setup_house, 2),
        "setup_customer_usd": round(setup_customer, 2),
        "breakdown": breakdown,
        "methods_used": sorted(methods_used),
    }


def process_document(file_bytes, file_extension, mode_override=None, garment=None):
    """End-to-end. Handles PDF, JPG, PNG, SVG. mode_override: 'document', 'logo', or None for auto."""
    ext = file_extension.lower().lstrip(".")

    if ext == "pdf":
        page_images = pdf_to_images(file_bytes)
        primary_image = page_images[0]
        raster_bytes = image_to_bytes(primary_image)
        page_count = len(page_images)
    elif ext == "svg":
        primary_image = svg_to_pil(file_bytes)
        raster_bytes = image_to_bytes(primary_image)
        page_count = 1
    else:
        primary_image = Image.open(BytesIO(file_bytes))
        raster_bytes = file_bytes
        page_count = 1

    if mode_override in ("document", "logo"):
        mode = mode_override
        vision_features = analyze_image_features(raster_bytes) if mode == "logo" else None
    else:
        mode, vision_features = auto_detect_mode(raster_bytes, ext)

    result = {
        "input_image": primary_image,
        "page_count": page_count,
        "mode": mode,
        "auto_detected": mode_override not in ("document", "logo"),
        "input_format": ext,
    }

    if mode == "logo":
        if vision_features is None:
            vision_features = analyze_image_features(raster_bytes)
        result["logo_analysis"] = vision_features
        if garment:
            result["logo_recommendations"] = get_logo_recommendations(vision_features, garment, input_format=ext)
    else:
        ocr_text = ocr_image_bytes(raster_bytes)
        result["raw_ocr"] = ocr_text
        result["structured"] = extract_structured(ocr_text) if ocr_text.strip() else None

    return result


# ============================================================
# UNIVERSAL TEXT EXTRACTOR (Module 7)
# Get text from anything, then have Claude tell you what it is.
# ============================================================

UNIVERSAL_CONTEXT_PROMPT = """You are an analyst reviewing text extracted from a file someone uploaded.

FILENAME: {filename}
EXTRACTION METHOD: {method}
TEXT LENGTH: {length} characters

EXTRACTED TEXT:
---
{text}
---

Provide a concise contextual analysis as plain prose (no markdown headers, no bullet lists - just a few short paragraphs):

1. What is this document? (1 sentence)
2. What is its likely purpose? (1-2 sentences)
3. The 3-5 most important pieces of information or data points in it.
4. Anything notable, unusual, or that stands out.

Keep it under 250 words total. Be specific to THIS content - not generic."""


def extract_text_universal(file_bytes, file_extension):
    """Extract text from any supported file type. Returns (text, method_used).

    Supports: pdf, ai, jpg, jpeg, png, webp, gif, bmp, tiff, svg, docx, txt, md, csv, log.
    """
    ext = file_extension.lower().lstrip(".")

    if ext in ("txt", "md", "csv", "log", "json", "xml", "html", "htm"):
        try:
            text = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            text = file_bytes.decode("latin-1", errors="replace")
        return text, f"Direct text decode (.{ext})"

    if ext == "docx":
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for tbl in doc.tables:
                for row in tbl.rows:
                    row_text = " | ".join(c.text.strip() for c in row.cells)
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts), "python-docx (paragraphs + tables)"
        except ImportError:
            return "", "python-docx not installed (pip install python-docx)"
        except Exception as e:
            return "", f"DOCX error: {e}"

    if ext in ("pdf", "ai"):
        try:
            page_images = pdf_to_images(file_bytes)
            page_texts = []
            for i, img in enumerate(page_images):
                page_bytes = image_to_bytes(img)
                page_text = ocr_image_bytes(page_bytes)
                if page_text.strip():
                    page_texts.append(f"--- Page {i+1} ---\n{page_text}")
            return "\n\n".join(page_texts), f"Cloud Vision OCR ({len(page_images)} pages)"
        except Exception as e:
            return "", f"PDF/AI processing error: {e}"

    if ext == "svg":
        try:
            img = svg_to_pil(file_bytes)
            page_bytes = image_to_bytes(img)
            text = ocr_image_bytes(page_bytes)
            return text, "SVG rasterize + Cloud Vision OCR"
        except Exception as e:
            return "", f"SVG processing error: {e}"

    if ext in ("jpg", "jpeg", "png", "webp", "gif", "bmp", "tiff", "tif"):
        try:
            text = ocr_image_bytes(file_bytes)
            return text, f"Cloud Vision OCR (.{ext})"
        except Exception as e:
            return "", f"Image OCR error: {e}"

    return "", f"Unsupported file type: .{ext} (try converting to PDF, image, or DOCX)"


def get_universal_context(extracted_text, filename, method):
    """Send extracted text to Claude for contextual analysis."""
    client = get_anthropic_client()
    max_text = 20000
    text_for_prompt = extracted_text[:max_text]
    if len(extracted_text) > max_text:
        text_for_prompt += f"\n\n[...truncated at {max_text} chars; full text was {len(extracted_text)} chars]"

    prompt = UNIVERSAL_CONTEXT_PROMPT.format(
        filename=filename,
        method=method,
        length=len(extracted_text),
        text=text_for_prompt or "(no text extracted)",
    )
    msg = client.messages.create(
        model=EXTRACTION_MODEL, max_tokens=800,
        messages=[{"role": "user", "content": prompt}],
    )
    analysis = msg.content[0].text.strip()
    return {
        "analysis": analysis,
        "_meta": {
            "model": msg.model,
            "tokens_in": msg.usage.input_tokens,
            "tokens_out": msg.usage.output_tokens,
            "approx_cost_usd": round((msg.usage.input_tokens * 0.0000008) + (msg.usage.output_tokens * 0.000004), 4),
        }
    }
